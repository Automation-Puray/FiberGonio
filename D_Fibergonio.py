#!/usr/bin/env python3
"""
DexArm UV grid-scan – SeaBreeze + HDF5
====================================================
* Full spectrum (3648 px) stored per point
* Data written to 8_EduardoScanData.h5:
    ├─ /spectra      [GRID_ROWS, GRID_COLUMNS, 3648]   float32
    ├─ /wavelengths  [3648]                            float32
    └─ file attrs    x_step_mm, y_step_mm, grid_origin
"""

import os
import json
import serial
import asyncio
import logging
import sys
import time
import datetime
import subprocess
import numpy as np
import h5py
import matplotlib.pyplot as plt
import seabreeze.spectrometers as sb
#import usbtmc
import pyvisa
import aiohttp
from aioairq import AirQ
from concurrent.futures import ThreadPoolExecutor
from scipy.interpolate import griddata
from math import nan, isnan
from pydexarm import Dexarm
from docx import Document
from datetime import datetime
from docx.shared import Inches
import glob
#TODO: Non Hard-Coded Grid

#TODO: Connection and control of Motorised Rotational Stages

#TODO: Visualization Optimization

#TODO: Append to file after each measurement to not overload RAM. Prepare for Bigger Data Packages

# ────────────────────────────────────────────────────────────────────────────
# Hard Coded
# ────────────────────────────────────────────────────────────────────────────

# LASER
SERIAL_PORT_LASER = '/dev/ttyUSB0'
BAUD_RATE_LASER = 9600
MAX_POWER_MW = 570
MACROS = {'on': '7E02010101', 'off': '7E02010100'}

# ARDUINO - Rotational Stage
SERIAL_PORT_ARDUINO = '/dev/ttyUSB1'
BAUD_RATE_ARDUINO = 115200

# SPECTROMETER
FULL_SCALE          = 65535                                                                                                                     # 16-bit ADC max for USB4000/Flame-T
SCANS_TO_AVERAGE    = 10
BOXSCAR             = 2
DEFAULT_INTEGRATION = 100000
T_INTEGRATION_MIN   = 1
T_INTEGRATION_MAX   = 10000000
SATURATION_LIMIT    = 60000
SUBRANGE_MIN        = 300
SUBRANGE_MAX        = 500
#CALIBRATION_PATH = 'path/to/file.irradcal'

# ROBOT
SERIAL_PORT_ROBOT         = "/dev/ttyACM0"
FEEDRATE_NORMAL     = 2000                                                                                                                      # mm / min
FEEDRATE_FINE       = 1000                                                                                                                      # mm / min
PARK_X              = 100.0                                                                                                                     # mm
PARK_Y              = 70.0                                                                                                                      # mm  --> Measurement Position
PARK_Z              = -80.0                                                                                                                     # mm
Z_HOMING_OFFSET     = 300.0                                                                                                                     # mm
STEP_SIZES          = [200.0, 100.0, 50.0, 20.0, 10.0, 5.0, 1.0]                                                                                #TODO: Steps 0.2 and 0.1 # mm
SYSTEM_BOUNDS = {'x_min': 0.0, 'x_max': 1000.0, 'y_min': -100.0, 'y_max': 120.0, 'z_min': -90.0, 'z_max': 60.0}
FIX_MIN_BOUND = 80.0
MARGIN = 3                                                                                                                                      # safety delay (s)
EXTRA_MARGIN = 2

#AIRQ 
AIRQ_ADDRESS = "10.42.0.125" # IP of the AirQ device
AIRQ_PASSWORD = "LightLabairQ"  # Password for the AirQ device

# ────────────────────────────────────────────────────────────────────────────
# Logging setup
# ────────────────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s  %(levelname)7s:  %(message)s", datefmt="%H:%M:%S")
measurement_time = time.strftime("%Y-%m-%d %H:%M:%S")




# ────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────

def unpack_xyz(raw):
    """DexArm returns (x, y, z, e, a, b, c) we need myX = E, myY = Z, myZ = Y  to imitate a vertical plane."""
    return raw[3], raw[2], raw[1]

def pack_xyz(x, y, z):
    return {"e": x, "y": z, "z": y}

def in_bounds(val, low, high):
    return isnan(low) or (low <= val <= high)

def get_key():
    """Cross-platform single-key input (arrow/WASD/enter/q)."""
    try:                                                                                                                                        # Windows
        import msvcrt
        if msvcrt.kbhit():
            ch = msvcrt.getch()
            if ch in (b"\x00", b"\xe0"):
                ch = msvcrt.getch()
                return {b"H": "UP", b"P": "DOWN", b"K": "LEFT", b"M": "RIGHT"}.get(ch, "")
            return "ENTER" if ch == b"\r" else ch.decode()
        return None
    except ImportError:                                                                                                                         # POSIX
        import tty, termios
        fd = sys.stdin.fileno()
        old = termios.tcgetattr(fd)
        try:
            tty.setraw(fd)
            c = sys.stdin.read(1)
            if c == "\x03":
                raise KeyboardInterrupt  # Manually raise if Ctrl+C pressed
            if c == "\x1b":
                seq = sys.stdin.read(2)
                return {"[A": "UP", "[B": "DOWN", "[C": "RIGHT", "[D": "LEFT"}.get(seq, "")
            return "ENTER" if c in ("\r", "\n") else c
        finally:
            termios.tcsetattr(fd, termios.TCSADRAIN, old)



def initialize_arduino():
    try:
        ser = serial.Serial(
            port=SERIAL_PORT_ARDUINO,
            baudrate=BAUD_RATE_ARDUINO,
            timeout=2
        )
        print(f"arduino: {SERIAL_PORT_ARDUINO} open")
        return ser
    except serial.SerialException as e:
        logging.error(f"✗  Error opening serial port {SERIAL_PORT_ARDUINO}: {e}")
        return None

def send_angle_to_arduino(ser, angle):
    if ser is None:
        logging.error("Arduino serial connection not available.")
        return
    try:
        command = f"{int(angle)}\n"
        ser.write(command.encode('utf-8'))
        #logging.info(f"→ Sent to Arduino: {angle}")
    except Exception as e:
        logging.error(f"✗ Failed to send to Arduino: {e}")

def initialize_serial():
    try:
        laser = serial.Serial(
            port = SERIAL_PORT_LASER,
            baudrate = BAUD_RATE_LASER,
            bytesize = serial.EIGHTBITS,
            parity = serial.PARITY_NONE,
            stopbits = serial.STOPBITS_ONE,
            timeout = 1
        )
        print(f"laser: {SERIAL_PORT_LASER} open")
        return laser
    except serial.SerialException as e:
        logging.error(f"✗  Error opening serial port {SERIAL_PORT_LASER}: {e}")
        exit(1)

def send_command(laser, hex_string, expect_response=True):
    try:
        data = bytes.fromhex(hex_string)
        laser.write(data)
        if False:                                                                                                                               #Make True for DEBUG
            response = laser.read(64) if expect_response else b''
            print(f" Sent:     {hex_string}")
            print(f" Received: {response.hex() if response else 'No response'}\n")
    except Exception as e:
        print(f" Error sending command: {e}")

def set_power(laser, power_mw):
    if not (0 <= power_mw <= MAX_POWER_MW):
        print(f" Power must be between 0 and {MAX_POWER_MW} mW.")
        return
    power_hex = f"{power_mw:04X}"  # Convert to 2-byte hex (big-endian)
    hex_string = f"7E040102{power_hex[:2]}{power_hex[2:]}"
    send_command(laser, hex_string)

def safe_laser_transition(laser):
    set_power(laser, 0)
    logging.info("✓ Laser calibrated and turned OFF for safety.")
    print("Please now disconnect the laser from the power meter and connect the laser to the fiber optic cable..")
    input("Press ENTER to confirm the fiber is connected and you are ready to proceed...")
    logging.info("proceeding with the experiment")
#calibration using pm400

def load_laser_calibration(json_path):
    with open(json_path, 'r') as f:
        data = json.load(f)
    outputs = np.array([entry['measured_output_mw'] for entry in data])
    setpoints = np.array([entry['setpoint_mw'] for entry in data])
    sort_idx = np.argsort(outputs)
    return setpoints[sort_idx], outputs[sort_idx]

def find_setpoint_for_output(desired_output_mw, setpoints, outputs):
    return float(np.interp(desired_output_mw, outputs, setpoints))

async def ramp_laser_to_setpoint(laser, target_setpoint, steps=5, delay=8):
    current = 0
    step_size = (target_setpoint - current) / steps
    for i in range(1, steps + 1):
        next_setpoint = current + step_size
        set_power(laser, int(round(next_setpoint)))
        #logging.info(f"  [Ramp] Step {i}/{steps}: Setpoint={next_setpoint:.1f} mW")
        await asyncio.sleep(delay)
        current = next_setpoint

async def feedback_set_laser_power(laser, inst, desired_output_mw, setpoints, outputs, tolerance=3, max_iter=12):
    logging.info("laser calibration and stablizing on progress...")
    setpoint = find_setpoint_for_output(desired_output_mw, setpoints, outputs)
    await ramp_laser_to_setpoint(laser, setpoint, steps=5, delay=15)
    
    history_setpoints = []
    history_measured = []
    
    
    for attempt in range(1, max_iter+1):
        set_power(laser, int(round(setpoint)))
        await asyncio.sleep(5 if attempt < max_iter - 2 else 10)
        measured_mw = get_pm400_power(inst)
        if measured_mw is None:
            logging.warning("PM400 read failed. Retrying...")
            continue
        #logging.info(f"Attempt {attempt}: Setpoint={setpoint:.2f} mW, Measured={measured_mw:.2f} mW")
        
        history_setpoints.append(setpoint)
        history_measured.append(measured_mw)
        
        
        error = desired_output_mw - measured_mw
        gain = 0.7 if abs(error) > 2 else 0.2

        if abs(error) <= tolerance:
            #logging.info(f"Within tolerance (±{tolerance} mW). Waiting extra 30s to confirm stability...")
            time.sleep(5)
            measured_mw2 = get_pm400_power(inst)
            #logging.info(f"Re-measured after extra delay: {measured_mw2:.2f} mW")
            error2 = desired_output_mw - measured_mw2

            
            history_setpoints.append(setpoint)
            history_measured.append(measured_mw2)


            if abs(error2) <= tolerance:
                #logging.info(f"Target achieved: {measured_mw2:.2f} mW (within ±{tolerance} mW) and stable after extra delay.")
                return setpoint, measured_mw2
            else:
                #logging.warning(f"Drift detected after extra delay (error now {error2:.2f} mW). Continuing feedback loop.")
                setpoint += error2 * gain
                setpoint = max(0, min(setpoint, MAX_POWER_MW))
                continue
        setpoint += error * gain
        setpoint = max(0, min(setpoint, MAX_POWER_MW))
    #logging.warning(f"Max attempts reached. Final measured output: {measured_mw:.2f} mW")
    return setpoint, measured_mw

def save_session(flame, custom_home, measurement_bounds):
    session_data = {
        "integration_time_us": flame.integration_time_us(),
        "custom_home": custom_home,
        "measurement_bounds": measurement_bounds,
        "dark_path": "dark_spectrum.npy",
        "calibration_path": "calibration.npy"
    }

    # Save metadata as JSON
    with open("session.json", "w") as f:
        json.dump(session_data, f, indent=2)

    # Save binary arrays
    if flame._dark is not None:
        np.save(session_data["dark_path"], flame._dark)
    if flame._calibration is not None:
        np.save(session_data["calibration_path"], flame._calibration)

    logging.info("✓ Session saved to session.json + .npy files")

def load_session():
    try:
        with open("session.json", "r") as f:
            session_data = json.load(f)

        integration_time_us = session_data.get("integration_time_us", DEFAULT_INTEGRATION)
        custom_home = session_data.get("custom_home", [0.0, 0.0, 0.0])
        measurement_bounds = session_data.get("measurement_bounds", SYSTEM_BOUNDS.copy())

        dark = np.load(session_data["dark_path"])
        calibration = np.load(session_data["calibration_path"])

        logging.info("✓  Session restored from session.json")

        return integration_time_us, custom_home, measurement_bounds, dark, calibration

    except Exception as e:
        logging.warning(f"✗ Failed to load session: {e}")
        return DEFAULT_INTEGRATION, [0.0, 0.0, 0.0], SYSTEM_BOUNDS.copy(), None, None

def clear_terminal():
    subprocess.run("clear")

def list_h5_files(directory="."):
    return [f for f in os.listdir(directory) if f.endswith(".h5")]

def select_h5_file(directory="."):
    files = list_h5_files(directory)
    if not files:
        print("No .h5 files found in the directory.")
        return None

    print("\nAvailable .h5 files:")
    for idx, fname in enumerate(files, start=1):
        print(f"  {idx}. {fname}")

    while True:
        try:
            choice = int(input("            INPUT:  Select a file number to load: "))
            if 1 <= choice <= len(files):
                return os.path.join(directory, files[choice - 1])
            else:
                print("Invalid number.")
        except ValueError:
            print("Please enter a valid integer.")
#PM400
def get_pm400_power(inst):
    try:
        power_w = float(inst.query("MEAS:POW?"))
        return power_w * 1000  # Return in mW
    except Exception as e:
        logging.error(f"PM400 power read failed: {e}")
        return None
#AirQ
async def get_env_conditions():
    #get the temprature and humidity - will add other paras if needed
    try:
        async with aiohttp.ClientSession() as session:
            airq = AirQ("10.42.0.125", "LightLabairQ", session)
            data = await airq.get_latest_data()
            return data.get('temperature'), data.get('humidity')
    except Exception as e:
        logging.error(f"AirQ data fetch failed: {e}")
        return None, None



#Document_Generation

def generate_measurement_report(h5_filename, template_filename, output_filename):
    try:
        # Read attributes from HDF5
        with h5py.File(h5_filename, "r") as h5f:
            base = os.path.splitext(os.path.basename(h5_filename))[0]
            
            replacements = { 
                #"{{pm400_power_w}}": str(h5f.attrs.get('pm400_power_w', 'N/A')),
                "{{airq_temperature_c}}": str(h5f.attrs.get('airq_temperature_c', 'N/A')),
                "{{airq_humidity_pct}}": str(h5f.attrs.get('airq_humidity_pct', 'N/A')),
                "{{fiber_name}}": str(h5f.attrs.get('fiber_name', 'N/A')),
                "{{measurement_time}}": str(h5f.attrs.get('measurement_time', 'N/A')),
                "{{date}}": datetime.now().strftime("%Y-%m-%d"),
                
                # will Add other paras once we defined the complete report structure
            }
        pattern = os.path.join(os.path.dirname(h5_filename), f"*int3DPLOT_{base}.png")
        matches = glob.glob(pattern)
        plot_filename = matches[0] if matches else None
        # Load the template
        doc = Document(template_filename)

        # Replace placeholders in paragraphs
        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, val)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, val)
        # Replace placeholders in headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for key, val in replacements.items():
                    if key in paragraph.text:
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, val)
            # Replace placeholders in header tables                    
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, val in replacements.items():
                                if key in paragraph.text:
                                    for run in paragraph.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, val)
        if plot_filename and os.path.exists(plot_filename):
            doc.add_paragraph("3D Intensity Plot:")
            doc.add_picture(plot_filename, width=Inches(6))  # adjust size as needed
        else:
            logging.warning(f"Plot image not found: {plot_filename}")
        # Save the filled document
        doc.save(output_filename)
        logging.info(f"Report generated: {output_filename}")

    except Exception as e:
        logging.error(f"Failed to generate report: {e}")



class AsyncFlame:
    """
    Non-blocking wrapper for a SeaBreeze-compatible spectrometer (e.g., Flame, USB4000).
    Includes integration time control, dark reference storage, spectral calibration,
    saturation detection, and post-measurement spectrum integration.
    """

    def __init__(self, integration_time_us: int = DEFAULT_INTEGRATION):
        self.integration_time = integration_time_us
        self._dev = None
        self._dark = None
        self._wavelengths = None
        self._calibration = None
        self._pool = ThreadPoolExecutor(max_workers=1)

    # ---------- blocking helpers (run in executor) -------------------------
    def _sync_open(self):
        device_ids = sb.list_devices()
        if not device_ids:
            raise RuntimeError("No spectrometer found via SeaBreeze")
        
        self._dev = sb.Spectrometer(device_ids[0])
        self._dev.integration_time_micros(self.integration_time)
        print(f"spectrometer: {device_ids[0]} open")
        for attr, arg in (("scans_to_average", SCANS_TO_AVERAGE), ("boxcar_width", BOXSCAR)):
            try:
                getattr(self._dev, attr)(arg)
            except AttributeError:
                pass

        self._dark = np.asarray(self._dev.intensities(), dtype=np.float32)
        self._wavelengths = np.asarray(self._dev.wavelengths(), dtype=np.float32)
        self._load_calibration()

    def _load_calibration(self, calib_path="9_FLMT09336_cc_20230822.IRRADCAL"):
        try:
            calib_data = np.loadtxt(calib_path, skiprows=9)
            calib_wl = calib_data[:, 0]
            calib_vals = calib_data[:, 1]
            self._calibration = np.interp(self._wavelengths, calib_wl, calib_vals)
            logging.info("\u2713  Calibration data loaded and interpolated.")
        except Exception as e:
            logging.warning(f"\u2717 Calibration file could not be parsed: {e}. Using flat calibration (ones). Results will not be accurate.")
            self._calibration = np.ones_like(self._wavelengths)

    def calibrate_spectrum(self, spectrum: np.ndarray) -> np.ndarray:
        if self._calibration is None:
            raise RuntimeError("Calibration not loaded.")
        if self._dark is None:
            raise RuntimeError("Dark spectrum not available.")
        if self._wavelengths is None:
            raise RuntimeError("Wavelengths not initialized.")

        # Step 1: Dark subtraction
        corrected = np.maximum(spectrum - self._dark, 0.0)
        # Step 2: Multiply by calibration factors (µJ/count/nm)
        calibrated = corrected * self._calibration
        # Step 3: Convert integration time from µs to s
        integration_time_s = self.integration_time / 1_000_000
        # Step 4: Divide by integration time and sensor area
        sensor_area_cm2 = 0.11946
        calibrated /= (integration_time_s * sensor_area_cm2)
        # Step 5: Divide by Δλ (nm), using mean spacing
        delta_lambda = np.mean(np.diff(self._wavelengths))
        calibrated /= delta_lambda

        return calibrated                                                                                                                       # Units: µW / cm² / nm

    def check_saturation(self, spectrum: np.ndarray, saturation_threshold: float = SATURATION_LIMIT) -> tuple[bool, float, float]:
        peak = float(np.max(spectrum))
        pct = 100.0 * peak / FULL_SCALE
        is_saturated = peak >= saturation_threshold
        return is_saturated, peak, pct

    async def capture_dark_reference(self):
        input("            INPUT:  Prepare to take Dark Spectrum, cover spectrometer and press ENTER")
        raw = await self.read_spectrum()
        self._dark = raw
        logging.info("\u2713  Dark spectrum recorded")

    def integrate_spectrum(self, spectrum: np.ndarray, lower_nm: float = SUBRANGE_MIN, upper_nm: float = SUBRANGE_MAX) -> float:
        if self._wavelengths is None:
            raise RuntimeError("Wavelength array is not initialized.")
        if self._wavelengths.shape != spectrum.shape:
            raise ValueError("Wavelength and spectrum arrays must have the same shape.")

        mask = (self._wavelengths >= lower_nm) & (self._wavelengths <= upper_nm)
        if not np.any(mask):
            raise ValueError(f"No wavelengths found in range {lower_nm}–{upper_nm} nm.")

        return float(np.trapezoid(spectrum[mask], x=self._wavelengths[mask]))

    def _sync_acquire(self) -> np.ndarray:
        if self._dev is None:
            raise RuntimeError("Spectrometer device not initialized. Call open() first.")
        return np.asarray(self._dev.intensities(), dtype=np.float32)

    def _sync_close(self):
        if self._dev:
            self._dev.close()

    # ---------- async API ---------------------------------------------------
    async def open(self):
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self._pool, self._sync_open)

    async def read_spectrum(self) -> np.ndarray:
        if self._dev is None:
            raise RuntimeError("Spectrometer not opened. Did you forget to call await open()?")
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(self._pool, self._sync_acquire)

    async def set_integration_time(self, us: int = DEFAULT_INTEGRATION):
        # Clamp integration time to device-supported range
        us = max(T_INTEGRATION_MIN, min(us, T_INTEGRATION_MAX))
        self.integration_time = us
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self._pool, self._dev.integration_time_micros, us)

    async def close(self):
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self._pool, self._sync_close)

    @property
    def wavelengths(self) -> np.ndarray:
        return self._wavelengths

    def integration_time_us(self) -> int:
        return self.integration_time

async def park_robot(dexarm, x, y, z):
# ────────────────────────────────────────────────────────────────────────────
# Safe Park:
# Function to safely park the robot 
# ────────────────────────────────────────────────────────────────────────────
    dz_t = (abs(z-PARK_Z) / FEEDRATE_NORMAL) * 60                                                                               # Distance / Movement Rate = Time in min * 60 = Time in s
    dy_t = (abs(y-PARK_Y) / FEEDRATE_NORMAL) * 60
    dx_t = (abs(x-PARK_X) / FEEDRATE_NORMAL) * 60
    z = PARK_Z
    y = PARK_Y
    x = PARK_X
    logging.info(f"Moving Robot to Safe Position: X={x:.1f}, Y={y:.1f}, Z={z:.1f}")
    dexarm.move_to(**pack_xyz(None, None, z), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dz_t)                                                                                                   # Optional delay for clean shutdown
    dexarm.move_to(**pack_xyz(None, y, None), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dy_t)     
    dexarm.move_to(**pack_xyz(x, None, None), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dx_t)

    return x, y, z

async def home_robot(dexarm, x, y, z, custom_home):
# ────────────────────────────────────────────────────────────────────────────
# Home:
# Function to safely park the robot 
# ────────────────────────────────────────────────────────────────────────────

    dz_t = (abs(z-custom_home[2]) / FEEDRATE_NORMAL) * 60 + MARGIN                                                                       # Distance / Movement Rate = Time in min * 60 = Time in s
    dy_t = (abs(y-custom_home[1]) / FEEDRATE_NORMAL) * 60 + MARGIN
    dx_t = (abs(x-custom_home[0]) / FEEDRATE_NORMAL) * 60 + MARGIN
    x = custom_home[0]
    y = custom_home[1]
    z = custom_home[2]
    logging.info(f"Moving Robot to Initial Position: X={x:.1f}, Y={y:.1f}, Z={z:.1f}")
    dexarm.move_to(**pack_xyz(None, None, z), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dz_t)                                                                                                   # Optional delay for clean shutdown
    dexarm.move_to(**pack_xyz(None, y, z), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dy_t)
    dexarm.move_to(**pack_xyz(x, None, None), feedrate=FEEDRATE_NORMAL)
    await asyncio.sleep(dx_t)

    return x, y, z, custom_home

async def manual_control(flame: AsyncFlame, dexarm: Dexarm, laser, arduino, step, step_idx, custom_home, x, y, z, measurement_bounds):
# ────────────────────────────────────────────────────────────────────────────
# Manual Jogging:
# Gives user manual keyboard control of the robot and lets him calibration the setup for the following measurements.
# Calibration of: Integration Time, Dark Spectrum, Measurement Bounds, Starting Point
# ────────────────────────────────────────────────────────────────────────────
    try:
        print("\n-- Use the controls to do the following: ")
        print("   1) Calibrate integration time & Dark Spectrum")
        print("   2) Set measurement boundaries")
        print("   3) Set Custom Home or the point where the measurement starts\n")

        print("-- Controls --")
        print("   ←/→ = X axis (DexArm E) – left/right")
        print("   ↑/↓ = Y axis (DexArm Z) – up/down")
        print("   W/S = Z axis (DexArm Y) – forward/backward\n")

        print("   M = Measure   |   P = Change laser power")
        print("   I = Set integration time and update dark spectrum\n")


        print("   R = Move Rotational Stages   |   G = Go to coordinate")
        print("   C = Change step size   |   X = Reset Measurement Boundaries")
        print("   H = Set new custom home (Meas Start Point and Max Boundary)\n")

        print("ENTER = continue   |   Ctrl + C = abort\n")

        logging.info(f"Current step size = {step:.1f} mm") 

        while True:
            k = (get_key() or "").upper()
            if not k:
                await asyncio.sleep(0.1)
                continue

            if k == "LEFT"      :
                if x + step <= measurement_bounds['x_max']:
                    x += step
                    dx_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dz_t = 0
                    dy_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (x upper bound = {measurement_bounds['x_max']})")
                    continue

            elif k == "RIGHT"   :
                if x - step >= measurement_bounds['x_min']:
                    x -= step
                    dx_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dz_t = 0
                    dy_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (x lower bound = {measurement_bounds['x_min']})")
                    continue

            elif k == "UP"      :
                if y + step <= measurement_bounds['y_max']:
                    y += step
                    dy_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dz_t = 0
                    dx_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (y upper bound = {measurement_bounds['y_max']})")
                    continue

            elif k == "DOWN"    :
                if y - step >= measurement_bounds['y_min']:
                    y -= step
                    dy_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dz_t = 0
                    dx_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (y lower bound = {measurement_bounds['y_min']})")
                    continue

            elif k == "W"       :
                if z + step <= measurement_bounds['z_max']:
                    z += step
                    dz_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dy_t = 0
                    dx_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (z upper bound = {measurement_bounds['z_max']})")
                    continue

            elif k == "S"       :
                if z - step >= measurement_bounds['z_min']:
                    z -= step
                    dz_t = (abs(step) / FEEDRATE_NORMAL) * 60                                                                                   # Distance / Movement Rate = Time in min * 60 = Time in s
                    dy_t = 0
                    dx_t = 0
                else:
                    logging.warning(f"✗  Robot limits reached (z lower bound = {measurement_bounds['z_min']})")
                    continue

            elif    k == "M"    :
                raw = await flame.read_spectrum()
                is_sat, peak, pct = flame.check_saturation(raw)
                cur_time = flame.integration_time_us()
                if is_sat:
                    logging.warning(f"✗  Saturation detected (Peak = {peak:.0f} counts, {pct:.1f} % of full scale)   |   Measured @ X = {x:.1f}, Y = {y:.1f}, Z = {z:.1f}   |   Current integration time = {cur_time} µs")
                else:
                    logging.info(f"✓  No saturation (Peak = {peak:.0f} counts, {pct:.1f} % of full scale)   |   Measured @ X = {x:.1f}, Y = {y:.1f}, Z = {z:.1f}   |   Current integration time = {cur_time} µs")
                continue

            elif    k == "P"        :
                try:
                    value = int(input(f"            INPUT:  Enter power (mW, 0-{MAX_POWER_MW}): ").strip())
                    if not 0 <= value <= MAX_POWER_MW:
                            raise ValueError
                    set_power(laser, value)
                    logging.info(f"✓  Power set to: {value} mW")
                except ValueError:
                    logging.warning(f"✗  Invalid input. Enter a number between 0 and {MAX_POWER_MW}.")
                    continue
                except (KeyboardInterrupt, asyncio.CancelledError):
                    raise       
                continue

            elif k == "I"       :
                set_power(laser, 0)
                cur_time = flame.integration_time_us()
                logging.info(f"Current integration time = {cur_time} µs")
                while True:
                    try:
                        raw = input("            INPUT:  Enter new integration time [µs] (1 – 10 000 000, or blank to cancel): ").strip()
                        if raw == "":
                            logging.info("↩  keeping current value.")
                            break
                        new_us = int(raw)
                        if not 1 <= new_us <= 10_000_000:
                            raise ValueError
                    except ValueError:
                        logging.warning("✗  Invalid input. Enter an integer between 1 and 10 000 000.")
                        continue
                    except KeyboardInterrupt:
                        raise                                                                                                                               # re-raise if you want to propagate up to the main loop
                    await flame.set_integration_time(new_us)
                    logging.info(f"✓  Integration time set to {new_us:,d} µs")
                    await flame.capture_dark_reference()
                    break
                continue

            elif    k == "R"        :
                try:
                    value = int(input(f"            INPUT:  Rotation in Degrees (-360° to 360°): ").strip())
                    if not -360 <= value <= 360:
                            raise ValueError
                    send_angle_to_arduino(arduino, value)                                                                                          # Send rotation angle
                    logging.info(f"✓  Rotational Stages moved: {value} °")
                except ValueError:
                    logging.warning(f"✗  Invalid input. Enter a number between -360° and 360°")
                    continue
                except (KeyboardInterrupt, asyncio.CancelledError):
                    raise       
                continue

            elif    k == "H"        :
                custom_home[0] = x
                custom_home[1] = y
                custom_home[2] = z
                logging.info(f"✓  Custom home set to: [{x:.2f}, {y:.2f}, {z:.2f}]")
                measurement_bounds['x_max'] = x
                measurement_bounds['z_max'] = z
                logging.info(f"✓  Measurement Boundaries updated: x_min = {measurement_bounds['x_min']:.2f}, x_max = {measurement_bounds['x_max']:.2f}, z_max = {measurement_bounds['z_max']:.2f} mm")
                continue

            elif    k == "G"        :
                while True:
                    try:
                        raw = input("            INPUT:  Enter coordinates where you want to move to separated by commas: X, Y, Z: ").strip()
                        if raw == "":
                            dx_t = 0
                            dy_t = 0
                            dz_t = 0
                            break
                        x_str, y_str, z_str = map(str.strip, raw.split(","))
                        x_new = float(x_str)
                        y_new = float(y_str)
                        z_new = float(z_str)
                        dz_t = (abs(z-z_new) / FEEDRATE_NORMAL) * 60                                                                            # Distance / Movement Rate = Time in min * 60 = Time in s
                        dy_t = (abs(y-y_new) / FEEDRATE_NORMAL) * 60                                                                            # Distance / Movement Rate = Time in min * 60 = Time in s
                        dx_t = (abs(x-x_new) / FEEDRATE_NORMAL) * 60                                                                            # Distance / Movement Rate = Time in min * 60 = Time in s
                        # Range check
                        if not (measurement_bounds['x_min'] <= x_new <= measurement_bounds['x_max']):
                            logging.warning(f"✗  X value out of bounds ({x_new:.1f} mm not in [{measurement_bounds['x_min']}, {measurement_bounds['x_max']}])")
                            dz_t = 0
                            dy_t = 0
                            dx_t = 0
                            continue
                        if not (measurement_bounds['y_min'] <= y_new <= measurement_bounds['y_max']):
                            logging.warning(f"✗  Y value out of bounds ({y_new:.1f} mm not in [{measurement_bounds['y_min']}, {measurement_bounds['y_max']}])")
                            dz_t = 0
                            dy_t = 0
                            dx_t = 0
                            continue
                        if not (measurement_bounds['z_min'] <= z_new <= measurement_bounds['z_max']):
                            logging.warning(f"✗  Z value out of bounds ({z_new:.1f} mm not in [{measurement_bounds['z_min']}, {measurement_bounds['z_max']}])")
                            dz_t = 0
                            dy_t = 0
                            dx_t = 0
                            continue
                        # Apply only if all checks pass
                        x, y, z = x_new, y_new, z_new

                    except ValueError:
                        logging.info("✗  Invalid input. Please enter 3 numeric values separated by commas (e.g., 80, 20, -100).")
                        continue
                    except (KeyboardInterrupt, asyncio.CancelledError):
                        raise
                    break

            elif    k == "C"        :
                step_idx = (step_idx + 1) % len(STEP_SIZES)
                step = STEP_SIZES[step_idx]
                logging.info(f"✓  Step size set to {step:.1f} mm")
                continue

            elif    k == "X"        :
                measurement_bounds = SYSTEM_BOUNDS.copy()
                measurement_bounds['x_min'] = FIX_MIN_BOUND
                logging.info(f"✓  Reset of Measurement Boundaries")
                continue

            elif    k == "ENTER"    :   break

            else:
                logging.warning("✗  Type in a valid command.")
                continue

            dexarm.move_to(**pack_xyz(None, None, z), feedrate=FEEDRATE_NORMAL)                                                                 # Move and Print Coordinates to the position
            await asyncio.sleep(dz_t)
            dexarm.move_to(**pack_xyz(None, y, None), feedrate=FEEDRATE_NORMAL)                                                                 # Move and Print Coordinates to the position
            await asyncio.sleep(dy_t)
            dexarm.move_to(**pack_xyz(x, None, None), feedrate=FEEDRATE_NORMAL)                                                                 # Move and Print Coordinates to the position
            await asyncio.sleep(dx_t)
            logging.info(f"Current Position: X = {x:.1f}, Y = {y:.1f}, Z = {z:.1f}")

        return custom_home, x, y, z, measurement_bounds
    except (KeyboardInterrupt, asyncio.CancelledError):
        raise                                                                                                                                   # Important: re-raise to propagate up
 
        
    



async def Measurement(flame: AsyncFlame, dexarm: Dexarm, laser, arduino, custom_home, x, y, z, measurement_bounds, grid_rows, grid_columns, spacing_x, spacing_y, wavelengths, num_pix, power_w, temperature , humidity,fiber_name, measurement_time):
# ────────────────────────────────────────────────────────────────────────────
# Measurement:
# Measurement along one line of the fiber optic cable. The cable or object is then rotated by a rotational motorised stage, and the measurement is repeated from another angle.
# ────────────────────────────────────────────────────────────────────────────
  
    try:
        # Prepare HDF5 Output File
        user_input = input("            INPUT:  Enter file name for HDF5 output (e.g., scan_data.h5): ").strip()
        if user_input.endswith(".h5"):
            user_input = user_input[:-3]  # remove .h5 for consistency

        # Add date prefix
        timestamp = time.strftime("%y%m%d")
        file_name = f"{timestamp}_{user_input}.h5"

        h5f = h5py.File(file_name, "w")

        #pm400 & airQ data stored in h5f 
        #h5f.attrs['pm400_power_w']      = power_w
        h5f.attrs['airq_temperature_c'] = temperature
        h5f.attrs['airq_humidity_pct']  = humidity
        h5f.attrs['fiber_name'] = fiber_name
        h5f.attrs['measurement_time'] = measurement_time
        

        spectra_dset = h5f.create_dataset(
            "spectra",
            shape=(grid_rows, grid_columns, num_pix),
            dtype="float32",
            chunks=(1, 1, num_pix),
            compression="gzip",
            fillvalue=np.nan,
        )
        h5f.create_dataset("wavelengths", data=wavelengths, dtype="float32")
        h5f.attrs.update(
            y_step_mm=spacing_y,
            x_step_mm=spacing_x,
            grid_origin=(0.0, 0.0),
        )

        # Prepare Variables
        dx_t = (spacing_x / FEEDRATE_FINE) * 60 + MARGIN
        total_points = grid_rows * grid_columns
        measured_points = 0

        

        # Coordinate buffers
        x_coord_map = np.zeros((grid_rows, grid_columns), dtype=np.float32)
        y_coord_map = np.zeros((grid_rows, grid_columns), dtype=np.float32)
        z_coord_map = np.zeros((grid_rows, grid_columns), dtype=np.float32)
        x_rep_map   = np.zeros((grid_rows, grid_columns), dtype=np.float32)
        y_ref_map   = np.zeros((grid_rows, grid_columns), dtype=np.float32)
        z_ref_map   = np.zeros((grid_rows, grid_columns), dtype=np.float32)

        z_offset = 1.0                                                                                                                                          # mm offset from object surface for z_ref

        for b in range(grid_rows):

            print("")                                                                                                                                                           #TODO: IF IN BOUNDS!!
            x, y, z, custom_home = await home_robot(dexarm, x, y, z, custom_home)

            angle_deg = b * spacing_y
            theta_rad = np.deg2rad(angle_deg)
            logging.info(f"Object Position:  = {angle_deg:.1f}°")
            send_angle_to_arduino(arduino, spacing_y)
            await asyncio.sleep(EXTRA_MARGIN)

            for c in range(grid_columns):
                x_rep = c * spacing_x
                x = custom_home[0] - x_rep

                if not in_bounds(x, measurement_bounds['x_min'], measurement_bounds['x_max']):                                                                                                  #TODO: IF IN BOUND DO OTHERWISE SAVE NULL AND GIVE WARNING, STILL SAVE COORDINATES
                    logging.warning(f"Measurement Ignored @ x = {x:.1f} mm, y = {y:.1f} mm, z = {z:.1f} mm ")
                    continue

                # Move
                dexarm.move_to(**pack_xyz(x, None, None), feedrate=FEEDRATE_FINE)
                await asyncio.sleep(dx_t)
                raw = await flame.read_spectrum()

                # Save robot-space coordinates
                x_coord_map[b, c] = x
                y_coord_map[b, c] = y
                z_coord_map[b, c] = z

                # Save reference-space coordinates (object-based)
                x_ref = x_rep
                y_ref = z_offset * np.cos(theta_rad)
                z_ref = -z_offset * np.sin(theta_rad)

                x_rep_map[b, c] = x_ref
                y_ref_map[b, c] = y_ref
                z_ref_map[b, c] = z_ref

                spectra_dset[b, c, :] = raw

                measured_points += 1
                progress = 100.0 * measured_points / total_points
                logging.info(f"Progress: {progress:.2f}% complete  →  Measured @ x = {x:.1f} mm, y = {y:.1f} mm, z = {z:.1f}  →  Corresponds to x' = {x_ref:.1f} mm, y' = {y_ref:.1f} mm, z' = {z_ref:.1f} mm")
                await asyncio.sleep(EXTRA_MARGIN)                                                                                                                                                     #wait too after the measurement has been done

        print("")
        send_angle_to_arduino(arduino, -360)
        await asyncio.sleep(EXTRA_MARGIN)
        x, y, z = await park_robot(dexarm, x, y, z)


        h5f.create_dataset("x_coord_map", data=x_coord_map)
        h5f.create_dataset("y_coord_map", data=y_coord_map)
        h5f.create_dataset("z_coord_map", data=z_coord_map)
        h5f.create_dataset("x_rep_map", data=x_rep_map)
        h5f.create_dataset("y_ref_map", data=y_ref_map)
        h5f.create_dataset("z_ref_map", data=z_ref_map)

        h5f.create_dataset("dark_spectrum", data=flame._dark)
        h5f.create_dataset("calibration_factors", data=flame._calibration)

        h5f.attrs.update(
            integration_time_us = flame.integration_time_us(),
            sensor_area_cm2 = 0.11946,
            x_step_mm = spacing_x,
            y_step_mm = spacing_y,
            grid_origin = (custom_home[0], custom_home[1], custom_home[2]),
            timestamp = time.strftime("%Y-%m-%d %H:%M:%S"),
            temperature_c = temperature,  
            humidity_pct = humidity 
        )

    except (KeyboardInterrupt, asyncio.CancelledError):
        raise

    finally:
        if h5f:
            h5f.close()
            logging.info("✓  HDF5 file closed.")
        set_power(laser, 0)
        logging.info("✓  Laser turned off")




# ────────────────────────────────────────────────────────────────────────────
# Main Routine
# ────────────────────────────────────────────────────────────────────────────
async def main():
    try:
        #clear_terminal()
        # ────────────────────────────────────────────────────────────────────────────
        # Setup
        # ────────────────────────────────────────────────────────────────────────────
        # Global Variables
        x_range   = 325.0                                                                                                                       # mm
        y_range   = 360.0                                                                                                                       # ° // This is the Angular Resolution - Not 0 when Rotation stages are there
        spacing_x = 25.0                                                                                                                        # mm per step
        spacing_y = 30.0                                                                                                                        # ° per step

        grid_columns = int(x_range / spacing_x) + 1                                                                                             # Measurement Plane Positions
        grid_rows    = int(y_range / spacing_y) + 1                                                                                             # Angular Positions

        step_idx = 3                                                                                                                            # Default Step Size = 20mm
        step = STEP_SIZES[step_idx]                                                                                                             # Current Step Size
        x=y=z = 0.0
        custom_home = [x,y,z]
        measurement_bounds = SYSTEM_BOUNDS.copy()
        measurement_bounds['x_min'] = FIX_MIN_BOUND
        
        fiber_name = input("Enter the fiber name: ")

        # Initialization
        #pm400               
        laser = initialize_serial()

        CALIB_JSON_PATH = '/home/Biomed/GonioMeasurements/laser_calibration.json'
        PM400_RESOURCE = 'USB0::4883::32885::P5005213::0::INSTR'
        TOLERANCE = 5  # mW

        # Load calibration
        calib_setpoints, calib_outputs = load_laser_calibration(CALIB_JSON_PATH)
        logging.info(" Laser calibration data loaded")
        desired_output = float(input("Enter desired output power in mW: "))

        # Initialize PM400 via pyvisa

        inst = None
        try:
            rm = pyvisa.ResourceManager()
            inst = rm.open_resource(PM400_RESOURCE)
            logging.info("✓ PM400 initialized: " + inst.query("*IDN?").strip())
        except Exception as e:
            logging.error(f"PM400 initialization failed: {e}")
            
        send_command(laser, MACROS["on"])
        await asyncio.sleep(0.1)

        if inst is not None:
            # PM400 available: use feedback loop
            final_setpoint, final_measured = await feedback_set_laser_power(
                laser, inst, desired_output, calib_setpoints, calib_outputs, tolerance=TOLERANCE, max_iter=5
            )
            power_w = final_measured
        else:
            # PM400 not available: fallback
            logging.warning("PM400 not available. Setting laser based on calibration only.")
            final_setpoint = find_setpoint_for_output(desired_output, calib_setpoints, calib_outputs)
            set_power(laser, int(round(final_setpoint)))
            power_w = None

        logging.info(f"✓ Laser turned ON and set to {final_setpoint:.2f} mW; expected output: {power_w if power_w is not None else 'unknown'} mW")
        safe_laser_transition(laser)
        set_power(laser, int(round(final_setpoint)))
        logging.info(f"Laser power restored to {final_setpoint:.2f} mW after fiber connection.")
        # Rotational Stages
        # Arduino
        arduino = initialize_arduino()
        await asyncio.sleep(0.1)

        # DexArm
        dexarm = Dexarm(port=SERIAL_PORT_ROBOT)
        await asyncio.sleep(2)

        # Spectrometer
        flame = AsyncFlame(integration_time_us = DEFAULT_INTEGRATION)                                                                           # 100 ms
        await flame.open()
        wavelengths = flame.wavelengths                                                                                                         # float32 array length 3648
        num_pix = len(wavelengths)
         
    
        def get_pm400_power(inst):
            try:
                power_w = float(inst.query("MEAS:POW?"))
                return power_w * 1000  # Return in mW
            except Exception as e:
                logging.error(f"PM400 power read failed: {e}")
                return None

        #AirQ

        temperature, humidity = await get_env_conditions()
        if temperature is not None and humidity is not None:
            logging.info(f"AirQ Temperature and Relative Humidity has been captured: {temperature:.1f}°C & {humidity:.1f}%")
        else:
            logging.info("AirQ Temperature and Relative Humidity not available")
            if temperature is None: temperature = float('nan')
            if humidity is None: humidity = float('nan')    

        # Reload Last Measurements Calibration
        int_time, custom_home, measurement_bounds, dark, calibration = load_session()
        await flame.set_integration_time(int_time)
        if dark is not None:
            flame._dark = dark
        if calibration is not None:
            flame._calibration = calibration

        if True:                                                                                                                                   # To be able to disable fast when evaluating.
            # Robot Position Init
            dexarm.sliding_rail_init()
            dexarm.go_home()
            dexarm.set_workorigin()
            x,y,z = list(unpack_xyz(dexarm.get_current_position()))
            z -=Z_HOMING_OFFSET                                                                                                                     # dirty hack, since dexarm has "his y = my forward" at 300 when homed

            # Moving Robot to Park Position
            x, y, z = await park_robot(dexarm, x, y, z)
        
        # ────────────────────────────────────────────────────────────────────────────
        # Loop
        # ────────────────────────────────────────────────────────────────────────────
        
        while True:
            print("")
            user_input = input("            INPUT:  System ready. Enter command: [measure (m), evaluate (e), quit (q)]: ").strip().lower()

            if user_input in ['measure', 'm']:

                try:

                    # Manual Jogging for Calibration Purposes
                    custom_home, x, y, z, measurement_bounds = await manual_control(flame, dexarm, laser, arduino, step, step_idx, custom_home, x, y, z, measurement_bounds)

                    try:
                        resp = input("            INPUT:  Do you want to save the measurement parameters (integration time, custom home, bounds)? (y/n): ").strip().lower()
                        if resp == 'y':
                            save_session(flame, custom_home, measurement_bounds)
                            logging.info("✓  Measurement parameters saved.")
                        else:
                            logging.info("↩  Measurement parameters not saved.")
                    except (KeyboardInterrupt, asyncio.CancelledError):
                        raise                                                                                                                   # Important: propagate the interrupt to the outer scope

                    # Measurement Routine
                    await asyncio.sleep(2)

                    


                    await Measurement(flame, dexarm, laser, arduino, custom_home, x, y, z, measurement_bounds, grid_rows, grid_columns, spacing_x, spacing_y, wavelengths, num_pix, power_w, temperature, humidity,fiber_name, measurement_time)
                    await asyncio.sleep(2)

                except (KeyboardInterrupt, asyncio.CancelledError):
                    logging.warning("↩  Aborted by user, returning to main menu.")
                    #Park Robot
                    await park_robot(dexarm, x, y, z)
                    #Turn Off Laser
                    logging.info("Turning laser off")
                    set_power(laser, 0)

                finally:
                    # Clean Up
                    logging.info("Session finished – DexArm parked")

            elif user_input in ['evaluate', 'e']:
                try:
                    file_path = select_h5_file(".")                                                                                             # Call file selector
                    if file_path is None:
                        raise FileNotFoundError("No .h5 file selected.")

                    # ── Load data ───────────────────────────────────────────────────────────────
                    with h5py.File(file_path, "r") as f:
                        wl = f["wavelengths"][:]
                        spectra = f["spectra"][:]
                        dark = f["dark_spectrum"][:]
                        calib_factors = f["calibration_factors"][:]

                        x_rep_map = f["x_rep_map"][:] if "x_rep_map" in f else None
                        y_ref_map = f["y_ref_map"][:] if "y_ref_map" in f else None
                        z_ref_map = f["z_ref_map"][:] if "z_ref_map" in f else None

                        x_step = f.attrs.get("x_step_mm", 100.0)
                        y_step = f.attrs.get("y_step_mm", 90.0)

                        integration_time_us = f.attrs.get("integration_time_us", 100_000)
                        sensor_area_cm2 = f.attrs.get("sensor_area_cm2", 0.11946)

                        print("")
                        logging.info("Metadata:")
                        for key, val in f.attrs.items():
                            logging.info(f"{key}: {val}")
                    
                    rows, cols, _ = spectra.shape

                    if False:
                    # ─────────────────────────────────────────────────────
                    # DEBUG PLOT – Raw, Dark, Calibrated Spectrum, and Calibration Factors
                    # ─────────────────────────────────────────────────────
                        debug_r, debug_c = 0, 1  # Grid position to check
                        raw = spectra[debug_r, debug_c, :]
                        dark_spectrum = dark
                        calib_interp = np.interp(wl, calib_wl, calib_factors)
                        
                        integration_time_s = integration_time_us / 1_000_000
                        delta_lambda = np.mean(np.diff(wl))
                        
                        calibrated = flame.calibrate_spectrum(raw)

                        # Plot
                        fig, axes = plt.subplots(1, 4, figsize=(22, 5), sharex=False)

                        # Dark spectrum
                        axes[0].plot(wl, dark_spectrum, color="black")
                        axes[0].set_title("Dark Spectrum")
                        axes[0].set_xlabel("Wavelength [nm]")
                        axes[0].set_ylabel("Counts")
                        axes[0].grid(True)

                        # Raw spectrum
                        axes[1].plot(wl, raw, color="gray")
                        axes[1].set_title(f"Raw Spectrum ({debug_r},{debug_c})")
                        axes[1].set_xlabel("Wavelength [nm]")
                        axes[1].set_ylabel("Counts")
                        axes[1].grid(True)

                        # Calibrated spectrum
                        axes[2].plot(wl, calibrated, color="blue")
                        axes[2].set_title("Calibrated Spectrum\n(Absolute Irradiance)")
                        axes[2].set_xlabel("Wavelength [nm]")
                        axes[2].set_ylabel("µW/cm²/nm")
                        axes[2].grid(True)

                        # Calibration curve
                        axes[3].plot(wl, calib_interp, color="green")
                        axes[3].set_title("Interpolated Calibration Factors")
                        axes[3].set_xlabel("Wavelength [nm]")
                        axes[3].set_ylabel("µJ / count / nm")
                        axes[3].grid(True)

                        # Adjust bottom margin for metadata
                        fig.subplots_adjust(bottom=0.2)

                        # Display parameters clearly below plots
                        param_text = (
                            f"Integration time: {integration_time_s:.6f} s     "
                            f"Sensor area: {sensor_area_cm2:.5f} cm²     "
                            f"Δλ (mean spacing): {delta_lambda:.4f} nm"
                        )
                        fig.text(0.5, 0.01, param_text, ha='center', fontsize=8)

                        plt.tight_layout()
                        plt.savefig("9_DebugEvaluationData.png", dpi=300)
                        print("✓ Saved diagnostic plot with calibration: 9_DebugEvaluationData.png")

                    if True:
                        # ─────────────────────────────────────────────────────
                        # PLOT – 4 Rows (Angles) × All Measured x′ Positions
                        # ─────────────────────────────────────────────────────

                        # Define desired angles (degrees) and map them to row indices
                        desired_angles = [0, 90, 180, 270]
                        rotation_step = y_step
                        angle_indices = [int(round(a / rotation_step)) for a in desired_angles]

                        # Extract all available x′ positions from map
                        x_grid_vals = x_rep_map[0, :]
                        x_indices = list(range(len(x_grid_vals)))
                        n_cols = len(x_indices)
                        n_rows = len(angle_indices)

                        # Initialize subplot grid
                        fig, axes = plt.subplots(n_rows, n_cols, figsize=(3 * n_cols, 2.5 * n_rows), sharex=True, sharey=True)
                        axes = np.atleast_2d(axes)

                        for row_idx, r_idx in enumerate(angle_indices):
                            for col_idx, c_idx in enumerate(x_indices):
                                ax = axes[row_idx, col_idx]

                                if r_idx >= spectra.shape[0] or c_idx >= spectra.shape[1]:
                                    ax.set_visible(False)
                                    continue

                                raw = spectra[r_idx, c_idx, :]
                                if np.isnan(raw).all():
                                    ax.set_visible(False)
                                    continue

                                # Calibrate and integrate
                                calibrated = flame.calibrate_spectrum(raw)
                                integral = flame.integrate_spectrum(calibrated)

                                # Extract coordinates
                                x_val = x_rep_map[r_idx, c_idx]
                                y_val = y_ref_map[r_idx, c_idx]
                                z_val = z_ref_map[r_idx, c_idx]

                                # Plot
                                ax.plot(wl, calibrated, color="green")
                                ax.set_title(
                                    f"θ={desired_angles[row_idx]}°\nx′={x_val:.0f} mm\nIrr={integral:.1f}",
                                    fontsize=7
                                )

                        # Axis labels
                        for ax in axes[-1, :]:
                            ax.set_xlabel("Wavelength [nm]")
                        for i in range(n_rows):
                            axes[i, 0].set_ylabel("µW/cm²/nm")

                        plt.tight_layout()

                        # Save plot
                        timestamp = time.strftime("%y%m%d")
                        base = os.path.splitext(os.path.basename(file_path))[0]
                        output_name = f"GRID_{base}_all_xs_fixed_angles.png"
                        plt.savefig(output_name, dpi=300)
                        logging.info(f"✓  Saved irradiance plot grid: {output_name}")
                        plt.close()


                    if True:
                    # ─────────────────────────────────────────────────────
                    # 3D PLOT Simple
                    # ─────────────────────────────────────────────────────

                        valid_indices = [(r, c) for r in range(rows) for c in range(cols) if not np.isnan(spectra[r, c, :]).all()]

                        # Load cylindrical coordinates (reference frame)
                        try:
                            x_grid = x_rep_map
                            y_grid = y_ref_map
                            z_grid = z_ref_map
                        except Exception as e:
                            logging.error(f"Failed to load coordinate maps: {e}")
                            exit(1)

                        # Compute irradiance per point
                        irr_vals = np.full((rows, cols), np.nan)
                        for r in range(rows):
                            for c in range(cols):
                                if not np.isnan(spectra[r, c, :]).all():
                                    calibrated = flame.calibrate_spectrum(spectra[r, c, :])
                                    irr_vals[r, c] = flame.integrate_spectrum(calibrated)

                        # ── Plotting ────────────────────────────────────────────────────────────────
                        fig = plt.figure(figsize=(12, 6))
                        ax = fig.add_subplot(111, projection='3d')

                        # Flatten original cylindrical grid and irradiance values
                        x_flat = x_grid.flatten()
                        y_flat = y_grid.flatten()
                        z_flat = z_grid.flatten()
                        irr_flat = irr_vals.flatten()

                        # Mask invalid entries
                        valid_mask = ~np.isnan(irr_flat)
                        points = np.column_stack((x_flat[valid_mask], z_flat[valid_mask], y_flat[valid_mask]))
                        values = irr_flat[valid_mask]

                        # Generate interpolated grid (cylinder surface)
                        x_lin = np.linspace(np.nanmin(x_grid), np.nanmax(x_grid), 200)
                        theta_lin = np.linspace(0, 2 * np.pi, 200)
                        X_grid, Theta_grid = np.meshgrid(x_lin, theta_lin)
                        Z_grid = 0.2 * np.cos(Theta_grid)  # radial slice Z
                        Y_grid = 0.2 * np.sin(Theta_grid)  # radial slice Y

                        interp_points = np.column_stack((X_grid.flatten(), Z_grid.flatten(), Y_grid.flatten()))
                        irr_interp = griddata(points, values, interp_points, method='linear', fill_value=np.nan)
                        irr_interp = irr_interp.reshape(X_grid.shape)

                        # Normalize for colormap
                        norm_irr = (irr_interp - np.nanmin(irr_interp)) / (np.nanmax(irr_interp) - np.nanmin(irr_interp))

                        # Plot interpolated irradiance surface
                        surf = ax.plot_surface(
                            X_grid, Z_grid, Y_grid,
                            facecolors=plt.cm.inferno(norm_irr),
                            rstride=1, cstride=1, linewidth=0, antialiased=False, shade=False
                        )

                        # Optional: fiber body
                        ax.plot_surface(X_grid, Z_grid, Y_grid, color="lightgray", alpha=0.05, linewidth=0, zorder=0)

                        # Dummy mappable for colorbar
                        mappable = plt.cm.ScalarMappable(cmap="inferno")
                        mappable.set_array(values)
                        cb = fig.colorbar(mappable, ax=ax, shrink=0.6, pad=0.1)
                        cb.set_label("Interpolated Irradiance [µW/cm²]")

                        # Axes setup
                        ax.set_xlabel("x' [mm]")
                        ax.set_ylabel("z' [mm]")
                        ax.set_zlabel("y' [mm]")
                        ax.set_title("Interpolated 3D Irradiance Surface")

                        ax.set_xlim(np.nanmin(x_grid), np.nanmax(x_grid))
                        ax.set_ylim(-0.3, 0.3)
                        ax.set_zlim(-0.3, 0.3)

                        # Save
                        
                        base = os.path.splitext(os.path.basename(file_path))[0]
                        output_name = f"int3DPLOT_{base}.png"
                        plt.tight_layout()
                        plt.savefig(output_name, dpi=300)
                        logging.info(f"✓  Saved interpolated 3D irradiance plot: {output_name}")
                        plt.close()

                    plt.close("all")  # Closes all open figures

                # Measurement Report Generation
                    if file_path:
                        report_name = os.path.basename(file_path).replace('.h5', '_report.docx')
                        output_path = os.path.join("Measurement_Report", report_name)
                        try:
                            generate_measurement_report(
                                h5_filename=file_path,
                                template_filename="gmf_temp.docx",
                                output_filename=output_path
                            )
                        except Exception as e:
                            logging.error(f"Could not generate report: {e}")
                except KeyboardInterrupt:
                    logging.warning("↩  Aborted by user, returning to main menu.")            

            elif user_input in ['quit', 'q']:

                # Clean Up
                await flame.close()                                                                                                        # Don't close, let open for EVAL
                print("")
                logging.info("✓  Flame spectrometer closed.")
                await asyncio.sleep(1)
                dexarm.close()
                logging.info("✓  DexArm connection closed.")
                await asyncio.sleep(1)
                if laser:
                    set_power(laser, 0)
                    laser.close()
                    logging.info("✓  Laser connection closed.")
                await asyncio.sleep(1)
                if arduino:
                    arduino.close()
                    logging.info("✓  Arduino connection closed.")
                await asyncio.sleep(1)
                break

            else:
                print("Invalid input. Options: measure (m), evaluate (e), quit (q).")

    except Exception as e:
        logging.error(f"An unexpected error occurred in main: {e}")


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n↩ Program interrupted by user (Ctrl+C). Exiting gracefully.")